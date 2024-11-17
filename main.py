from docx import Document
from functions import heading

# Load the sjabloon
template = "documents/template.docx"
document = Document(template)

#######################
#      Information    #
#######################

heading("Informatie")

# Ask for document name
docName = None
while not docName or " " in docName:
    docName = input("Geef een document naam op: ")

# Ask if solo
soloInput = input("Is dit een individueel project? (Ja/Nee): ")

while soloInput not in ["ja", "nee"]:
    soloInput = input("Ongeldige invoer. Is dit een individueel project? (Ja/Nee): ").lower()

solo = soloInput == "ja"

names = []

# Ask for name
if solo:
    name = input("Vul hier jouw volledige naam in: ")
else:
    while True:
        name = input("Voer een naam in (of type 'stop' om te stoppen): ")
    
        if name.lower() == "stop":
            break
    
        names.append(name)
        if names:
            formattedNames = ", ".join(names[:-1]) + " & " + names[-1] if len(names) > 1 else names[0]

# Ask for class
classNmr = input("Vul hier jouw klas in: ")

#######################
#       Voorblad      #
#######################

heading("Voorblad informatie")

# Add a title
title = document.add_paragraph(input("Vul hier een titel in: "), style="Title")

# Add the subject
subtitle = document.add_paragraph(input("Vul hier in voor welk vak het is: "), style="Subtitle")

# Add the author
author_date = document.add_paragraph()
if solo:
    author_date.add_run(f"Auteur: {name}\n")
else:
    author_date.add_run(f"Auteurs: {formattedNames}\n")

# Add the date
author_date.add_run(f"Datum: {input('Vul hier de datum in: ')}")

# Add the class number
author_date.add_run(f"Klas: {classNmr}")

# All of the next content will be on the next page
document.add_page_break()

#######################
#        Footer       #
#######################

footerSection = document.sections[0]

footer = footerSection.footer

# Ask for user's full name
footerParagraph0 = footer.paragraphs[0]
footerParagraph0.text = name if solo else formattedNames
footerParagraph0.style = 'Footer'

# Ask for class name
footerParagraph1 = footer.add_paragraph()
footerParagraph1.text = classNmr
footerParagraph1.style = 'Footer'

#######################
#       Content       #
#######################

heading(info="Table content info", width=45)

while True:
    date = input("Vul hier een datum in (of type 'stop' om te stoppen): ")
    if date.lower() == "stop":
        break

    # Add empty line 
    document.add_paragraph("\n", style="Normal")

    # Add the date as a heading
    document.add_paragraph(date, style="Heading 2")

    # Mak the table for the day
    table = document.add_table(rows=1, cols=3)
    table.style = "Logboek Style"

    # Ask for what happend that day
    while True:
        startTime = None
        while startTime is None or startTime == "":
            startTime = input("Wanneer ben je begonnen?: ")

        endTime = None
        while endTime is None or endTime == "":
            endTime = input("Wanneer ben je gestopt?: ")

        description = None
        while description is None or description == "":
            description = input("Wat heb je precies gedaan?: ")

        details = input("Waren er bijzonderheden?: ")

        # Add the table content
        from functions import addTableContent
        addTableContent(table, startTime, endTime, description, details)

        # Ask if the user want to add more days
        more = input("Wil je nog een activiteit toevoegen voor deze dag? (ja/nee): ")
        if more.lower() != "ja":
            break

# save the document
document.save(f"documents/{docName}.docx")
print(f"Document '{docName}' is succesvol gemaakt.")
